from flask import Flask, render_template, request, jsonify
import spacy
from spacy import displacy
import fitz  # PyMuPDF for PDF extraction
from docx import Document
from docx.shared import RGBColor
from io import BytesIO

app = Flask(__name__)
nlp = spacy.load("./output/model-best")

# Define color schemes for visualization and DOCX formatting
COLORS = {
    "PERSON": "lime",
    "LIT_WORK": "blue",
    "ART_WORK": "green",
    "ART_MOVEMENT": "orange",
    "ORG": "purple",
    "PLACE": "teal",
    "EVENT": "cyan",
    "GENRE": "pink",
    "CHARACTER": "brown",
    "QUOTE": "gold",
    "AWARD": "lime",
    "PERIOD": "magenta",
    "TECHNIQUE": "indigo",
    "MOVIE_TV": "violet"
}

ENTITY_COLORS = {
    "PERSON": RGBColor(50, 205, 50),       # Lime
    "LIT_WORK": RGBColor(0, 0, 255),         # Blue
    "ART_WORK": RGBColor(0, 128, 0),         # Green
    "ART_MOVEMENT": RGBColor(255, 165, 0),     # Orange
    "ORG": RGBColor(128, 0, 128),            # Purple
    "PLACE": RGBColor(0, 128, 128),          # Teal
    "EVENT": RGBColor(0, 255, 255),          # Cyan
    "GENRE": RGBColor(255, 20, 147),         # Pink
    "CHARACTER": RGBColor(165, 42, 42),      # Brown
    "QUOTE": RGBColor(255, 215, 0),          # Gold
    "AWARD": RGBColor(50, 205, 50),          # Lime
    "PERIOD": RGBColor(255, 0, 255),         # Magenta
    "TECHNIQUE": RGBColor(75, 0, 130),       # Indigo
    "MOVIE_TV": RGBColor(148, 0, 211)        # Violet
}

# Get the entity labels from the model's NER pipe and sort them.
MODEL_ENTITIES = sorted(nlp.get_pipe("ner").labels)


@app.route("/")
def index():
    """
    Render the main page with an initial (empty) visualization.
    """
    entity_types = MODEL_ENTITIES
    # Process an empty string for initial render
    doc = nlp("")
    options = {
        "ents": entity_types,
        "colors": {ent: COLORS.get(ent, "gray") for ent in entity_types}
    }
    html = displacy.render(doc, style="ent", options=options, page=False)
    return render_template("index.html", entity_types=entity_types, html=html)


@app.route("/filter", methods=["POST"])
def filter_entities():
    """
    Filter and highlight entities based on the selected types.
    """
    data = request.json
    selected_entities = data.get("selected_entities", [])
    user_text = data.get("text", "").strip()

    doc = nlp(user_text)
    options = {
        "ents": selected_entities,
        "colors": {ent: COLORS.get(ent, "gray") for ent in selected_entities}
    }
    html = displacy.render(doc, style="ent", options=options, page=False)
    return jsonify({"html": html})


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file given its byte content.
    """
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text("text") + "\n"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file given its byte content.
    """
    doc = Document(BytesIO(file_bytes))
    text = "\n".join(para.text for para in doc.paragraphs)
    return text.strip()


@app.route("/upload", methods=["POST"])
def upload_file():
    """
    Handle file uploads and extract text from PDF or DOCX files.
    """
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    filename = file.filename.lower()
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        return jsonify({"error": "Invalid file format. Please upload a PDF or DOCX."}), 400

    file_bytes = file.read()
    if filename.endswith(".pdf"):
        extracted_text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        extracted_text = extract_text_from_docx(file_bytes)

    return jsonify({"text": extracted_text})


@app.route("/save", methods=["POST"])
def save_results():
    """
    Save the processed text with highlighted entities to a DOCX file.
    """
    data = request.json
    text = data.get("text", "").strip()
    selected_entities = data.get("selected_entities", [])

    if not text:
        return jsonify({"error": "No content to save"}), 400

    doc = Document()
    para = doc.add_paragraph()
    processed_doc = nlp(text)

    last_index = 0
    for ent in processed_doc.ents:
        if ent.label_ in selected_entities:
            # Add unformatted text before the entity
            if ent.start_char > last_index:
                para.add_run(text[last_index:ent.start_char])
            # Add the entity text with bold and color formatting
            entity_run = para.add_run(ent.text)
            entity_run.bold = True
            entity_run.font.color.rgb = ENTITY_COLORS.get(
                ent.label_, RGBColor(128, 128, 128)
            )
            last_index = ent.end_char

    # Add any remaining text after the last entity
    if last_index < len(text):
        para.add_run(text[last_index:])

    file_path = "labeled_text.docx"
    doc.save(file_path)
    return jsonify({"success": True, "file_path": file_path})


if __name__ == "__main__":
    app.run(debug=True)
